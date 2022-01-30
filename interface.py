import tkinter as tk
import ctypes
from tkinter import filedialog
from PIL import ImageTk, Image
from tika import parser
from pytesseract import pytesseract
from docx import Document

# --------------- CONSTANTS ---------------
FONT = ("Hex Color Font", 11, "bold")
FG = "#111"
FG_BUTTON = "#fff"
BG = "#eee"
BG_BUTTON = "#518ff5"
ACTIVE_BG = "#4880DC"
ACTIVE_FG = "#fff"
LANGUAGES = ["Czech", "English", "Russian"]
LANGS_DICT = {"Czech": "ces", "English": "eng", "Russian": "rus"}


# --------------- CLASSES --------------
# Define a custom class with hover animation for all buttons
class DocButton(tk.Button):

    def __init__(self, **kw):
        super().__init__(**kw)
        self.bind("<Enter>", self.on_hover)
        self.bind("<Leave>", self.on_leave)
        super().config(
            width=10,
            fg=FG_BUTTON,
            bg=BG_BUTTON,
            font=FONT,
            relief="flat",
            highlightbackground=BG_BUTTON
        )

    def on_hover(self, event):
        if event:
            super().config(
                bg=ACTIVE_BG,
                fg=ACTIVE_FG
            )

    def on_leave(self, event):
        if event:
            super().config(
                bg=BG_BUTTON,
                fg=FG_BUTTON
            )


# Define a custom class for all labels
class DocLabel(tk.Label):
    def __init__(self, **kw):
        super().__init__(**kw)
        super().config(
            font=FONT,
            fg=FG,
            bg=BG
        )


class Application(tk.Tk):

    def __init__(self, master=None):

        # Config Root
        tk.Tk.__init__(self, master)
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
        self.geometry("800x600")
        self.tk.call('tk', 'scaling', 1.7)
        self.config(padx=25, pady=25, bg=BG)

        # Logo
        self.canvas = tk.Canvas(
            bg=BG,
            borderwidth=0,
            highlightthickness=0,
            width=128,
            height=128
        )
        self.logo_img = tk.PhotoImage(file="logo.png")
        self.canvas.create_image(64, 64, image=self.logo_img)
        self.canvas.grid(row=1, column=1, columnspan=3, padx=25, pady=25)

        # Main Label
        self.label_main = DocLabel(
            text="Browse Image or PDF and Convert to DOCX"
        )
        self.label_main.grid(row=2, column=1, columnspan=3, padx=25, pady=25)

        # Select Language Label and Dropdown Menu
        self.label_language = DocLabel(
            text="Select source language:"
        )
        self.label_language.grid(row=3, column=1, columnspan=2, padx=25, pady=25)
        self.lang_select = tk.StringVar(self)
        self.lang_select.set(LANGUAGES[0])
        self.arrow_img = ImageTk.PhotoImage(image=Image.open("arrow.png"))
        self.lang_menu = tk.OptionMenu(self, self.lang_select, *LANGUAGES)
        self.lang_menu.config(
            width=110,
            height=30,
            font=FONT,
            fg=FG_BUTTON,
            bg=BG_BUTTON,
            relief="flat",
            borderwidth=0,
            activebackground=ACTIVE_BG,
            activeforeground=ACTIVE_FG,
            indicatoron=0,
            image=self.arrow_img,
            compound="right"
        )
        self.lang_menu.image = self.arrow_img
        self.lang_menu.grid(row=3, column=3, columnspan=1)

        # Note Label
        self.label_note = DocLabel(text="")
        self.label_note.grid(row=4, column=1, columnspan=3)

        # Browse Button
        self.browse_button = DocButton(
            text='Browse',
            command=self.browse_files
        )
        self.browse_button.grid(row=5, column=1)

        # Convert Button
        self.convert_button = DocButton(
            text='Convert',
            command=self.recognise_and_convert
        )
        self.convert_button.grid(row=5, column=3)

        # Grid Config
        self.grid_rowconfigure(0, weight=1)
        self.grid_rowconfigure(5, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(4, weight=1)

        # Other Attributes
        self.file_path = ""
        self.target_path = ""

    # --------------- APP METHODS --------------
    def browse_files(self):
        filename = filedialog.askopenfilename(
            initialdir="/",
            title="Select a File",
            filetypes=(
                ("PDF or image files",
                 "*.pdf* *.jpeg *.jfif *.webp *.jpg *.png *.jpe *.svg *.bmp *.tif *.tiff *.gif"),
                ("all files",
                 "*.*")
            ))
        if filename:
            self.file_path = str(filename)
            self.target_path = f"{self.file_path.split('.')[0]}.docx"
            self.label_note.configure(text="File Opened: " + self.file_path.split("/")[-1], fg=FG)

    def read_img_text(self):
        path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        img = Image.open(self.file_path)
        pytesseract.tesseract_cmd = path_to_tesseract
        lang = LANGS_DICT[self.lang_select.get()]
        text = pytesseract.image_to_string(img, lang=lang)
        return text[:-1]

    def convert_img_text_to_docx(self):
        doc = Document()
        text = self.read_img_text()
        doc.add_paragraph(text)
        doc.save(self.target_path)
        self.label_note.configure(text="Image text converted.", fg="green")

    def convert_pdf_to_docx(self):
        raw = parser.from_file(self.file_path)
        text = raw['content']
        doc = Document()
        doc.add_paragraph(text)
        doc.save(self.target_path)
        self.label_note.configure(text="PDF converted.", fg="green")

    def recognise_and_convert(self):
        img_extensions = ["jpeg", "jfif", "webp", "jpg", "png", "jpe", "svg", "bmp", "tif", "tiff", "gif"]
        if self.file_path.split(".")[-1] == "pdf":
            self.convert_pdf_to_docx()
        else:
            if self.file_path.split(".")[-1] in img_extensions:
                self.convert_img_text_to_docx()
            else:
                self.label_note.configure(text="Wrong file format.", fg="red")
